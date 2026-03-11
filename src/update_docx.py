import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    # We create a new paragraph attached to the parent of the paragraph
    new_p = paragraph._parent.add_paragraph(text, style)
    # Then we move its xml element right after the current paragraph's xml element
    paragraph._p.addnext(new_p._p)
    return new_p

def main():
    base_dir = Path(__file__).resolve().parent.parent
    input_docx = base_dir / 'new_reviews' / 'test.docx'
    
    print(f"Opening {input_docx}")
    doc = Document(input_docx)
    
    insertions = {
        "平台型系统在中高学段应用更广泛": [
            (base_dir / 'output' / 'figures' / 'fig_a02_stage_waffle.png', "图2 学段分布华夫饼图"),
            (base_dir / 'output' / 'figures' / 'fig_a03_subject_lollipop.png', "图3 学科渗透棒棒糖图")
        ],
        "西部地区多数为零散单点分布，局部创新显著但整体规模有限": [
            (base_dir / 'output' / 'figures' / 'fig_a01_province_map.png', "图1 案例省域分布地图")
        ],
        "助研类智能产品尚未成为教育产品生态的重点方向，占比最少": [
            (base_dir / 'output' / 'figures' / 'fig_a04_scenario_treemap.png', "图4 应用场景树图")
        ]
    }
    
    # We iterate over a snapshot of the paragraphs so we don't mess up iteration when inserting
    paragraphs = list(doc.paragraphs)
    
    for p in paragraphs:
        text = p.text.strip()
        for key in list(insertions.keys()):
            if key in text:
                print(f"Found target paragraph for: {key[:20]}...")
                images = insertions[key]
                current_p = p
                for img_path, caption in images:
                    if img_path.exists():
                        # Paragraph for image
                        img_p = insert_after(current_p)
                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_p.add_run()
                        run.add_picture(str(img_path), width=Inches(5.5))
                        current_p = img_p
                        
                        # Paragraph for caption
                        cap_p = insert_after(current_p)
                        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap_run = cap_p.add_run(caption)
                        cap_run.font.size = Pt(10)
                        cap_run.font.bold = True
                        current_p = cap_p
                        
                        print(f"  Inserted {caption}")
                    else:
                        print(f"  [ERROR] Image not found: {img_path}")
                
                del insertions[key]
                break
                
    doc.save(input_docx)
    print(f"Updated {input_docx}")

if __name__ == "__main__":
    main()

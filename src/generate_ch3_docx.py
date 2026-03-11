import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def main():
    base_dir = Path(__file__).resolve().parent.parent
    output_docx = base_dir / 'Section_3_应用现状_最终版.docx'
    
    print(f"Creating {output_docx}")
    doc = Document()
    
    # Font settings
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    
    # Title
    doc.add_heading('第三章 人工智能赋能不同场景的应用现状', level=1)
    
    # 3.0
    doc.add_heading('3.0 总体场景分布与应用结构特征概述', level=2)
    
    p = doc.add_paragraph()
    p.add_run('当前人工智能在基础教育领域的应用已全面铺开，但不同场景的渗透深度呈现出高度不对称的结构性特征。基于收集的 1690 个典型案例，本报告将应用场景划分为“助学”、“助教”、“助评”、“助育”、“助管”与“助研”六大维度。通过对各维度案例体量的量化分析发现：')
    run = p.add_run('AI 教育应用展现出极强的“助学主导”与“助教协同”特征，而面向管理、评价与教研的系统级应用仍处于初期探索阶段。')
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run('具体而言，')
    run = p.add_run('“助学”场景')
    run.bold = True
    p.add_run('以压倒性的 1215 个案例（占比 71.89%）占据绝对核心地位，其中以“智能辅导系统”（958例）和“情境式学习”（181例）为代表的直接面向学生个体的个性化学习支持已成为最成熟的应用赛道。其次为')
    run = p.add_run('“助教”场景')
    run.bold = True
    p.add_run('（301例，占比 17.81%），AI 技术正快速向教师侧的“教学分析”（176例）与“备课/课件生成”（41例）环节渗透，助力教学质量的提升与教师减负。')
    
    p = doc.add_paragraph()
    p.add_run('相比之下，')
    run = p.add_run('“助育”')
    run.bold = True
    p.add_run('（58例，3.43%）、')
    run = p.add_run('“助评”')
    run.bold = True
    p.add_run('（63例，3.73%）、')
    run = p.add_run('“助管”')
    run.bold = True
    p.add_run('（39例，2.31%）与')
    run = p.add_run('“助研”')
    run.bold = True
    p.add_run('（8例，0.47%）场景由于涉及更深层次的教育伦理、多维度数据打通及复杂的专业体系重构，应用规模普遍较小。在这些长尾场景中，“综合素质评价”（45例）与“智能心理支持”（31例）表现出一定的创新潜力，代表了未来 AI 赋能教育向深水区迈进与全人发展的战略方向。')
    
    p = doc.add_paragraph()
    p.add_run('为了更清晰地呈现各个应用维度的结构，下表总括了人工智能应用场景在基础教育中的分布矩阵：')
    
    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '一级应用场景'
    hdr_cells[1].text = '总案例数'
    hdr_cells[2].text = '占比'
    hdr_cells[3].text = '核心二级应用场景分布提取'
    
    data = [
        ('助学', '1215', '71.89%', '智能辅导系统 (958例)，情境式学习 (181例)，游戏化学习 (19例)'),
        ('助教', '301', '17.81%', '教学分析 (176例)，教师备课 (41例)，作业管理 (25例)，课堂管理 (16例)'),
        ('助评', '63', '3.73%', '综合素质评价 (45例)，学生评估 (9例)，五育融合学生画像 (1例)'),
        ('助育', '58', '3.43%', '智能心理支持 (31例)，智能美育教育 (9例)，智能阅读辅助 (6例)，智能体育健康 (6例)'),
        ('助管', '39', '2.31%', '校园安全智能监控 (21例)，学生信息智能管理 (13例)，教务管理智能化 (3例)'),
        ('助研', '8', '0.47%', '智能教师专业发展平台 (6例)，智能科研助手 (1例)')
    ]
    
    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[0].paragraphs[0].runs[0].bold = True
        
        row_cells[1].text = item[1]
        row_cells[1].paragraphs[0].runs[0].bold = True

        row_cells[2].text = item[2]
        row_cells[2].paragraphs[0].runs[0].bold = True

        row_cells[3].text = item[3]
        
    doc.save(output_docx)
    print(f"Standalone docx generated at: {output_docx}")

if __name__ == "__main__":
    main()
